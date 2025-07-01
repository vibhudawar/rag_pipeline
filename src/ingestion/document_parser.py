from abc import ABC, abstractmethod
from typing import Dict, Any, List
import os
from pathlib import Path
import PyPDF2
import docx
from io import BytesIO


class DocumentParser(ABC):
    """Abstract base class for document parsers"""
    
    @abstractmethod
    def parse(self, file_path_or_bytes: Any) -> Dict[str, Any]:
        """Parse document and return text with metadata"""
        pass


class PDFParser(DocumentParser):
    """PDF document parser"""
    
    def parse(self, file_path_or_bytes: Any) -> Dict[str, Any]:
        """Parse PDF document"""
        try:
            if isinstance(file_path_or_bytes, (str, Path)):
                with open(file_path_or_bytes, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text() + "\n"
                    
                    return {
                        'text': text.strip(),
                        'metadata': {
                            'file_type': 'pdf',
                            'num_pages': len(pdf_reader.pages)
                        }
                    }
            else:  # Assume bytes-like object
                pdf_reader = PyPDF2.PdfReader(BytesIO(file_path_or_bytes))
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                
                return {
                    'text': text.strip(),
                    'metadata': {
                        'filename': 'uploaded_file.pdf',
                        'file_type': 'pdf',
                        'num_pages': len(pdf_reader.pages)
                    }
                }
        except Exception as e:
            raise RuntimeError(f"Failed to parse PDF: {str(e)}")


class DOCXParser(DocumentParser):
    """DOCX document parser"""
    
    def parse(self, file_path_or_bytes: Any) -> Dict[str, Any]:
        """Parse DOCX document"""
        try:
            if isinstance(file_path_or_bytes, (str, Path)):
                doc = docx.Document(file_path_or_bytes)
                filename = os.path.basename(file_path_or_bytes)
            else:  # Assume bytes-like object
                doc = docx.Document(BytesIO(file_path_or_bytes))
                filename = 'uploaded_file.docx'
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            return {
                'text': text.strip(),
                'metadata': {
                    'filename': filename,
                    'file_type': 'docx',
                    'num_paragraphs': len(doc.paragraphs)
                }
            }
        except Exception as e:
            raise RuntimeError(f"Failed to parse DOCX: {str(e)}")


class TXTParser(DocumentParser):
    """Plain text document parser"""
    
    def parse(self, file_path_or_bytes: Any) -> Dict[str, Any]:
        """Parse text document"""
        try:
            if isinstance(file_path_or_bytes, (str, Path)):
                with open(file_path_or_bytes, 'r', encoding='utf-8') as file:
                    text = file.read()
                    filename = os.path.basename(file_path_or_bytes)
            else:  # Assume bytes-like object
                text = file_path_or_bytes.decode('utf-8')
                filename = 'uploaded_file.txt'
            
            return {
                'text': text.strip(),
                'metadata': {
                    'filename': filename,
                    'file_type': 'txt',
                    'char_count': len(text)
                }
            }
        except Exception as e:
            raise RuntimeError(f"Failed to parse text file: {str(e)}")


class DocumentParserFactory:
    """Factory for creating appropriate document parsers"""
    
    _parsers = {
        '.pdf': PDFParser,
        '.docx': DOCXParser,
        '.txt': TXTParser,
        '.md': TXTParser  # Markdown files can use text parser
    }
    
    @classmethod
    def get_parser(cls, file_extension: str) -> DocumentParser:
        """Get appropriate parser for file extension"""
        parser_class = cls._parsers.get(file_extension.lower())
        if not parser_class:
            raise ValueError(f"Unsupported file type: {file_extension}")
        return parser_class()
    
    @classmethod
    def supported_extensions(cls) -> List[str]:
        """Get list of supported file extensions"""
        return list(cls._parsers.keys())


def parse_document(file_path_or_bytes: Any, file_extension: str = None) -> Dict[str, Any]:
    """Convenience function to parse any supported document"""
    if file_extension is None and isinstance(file_path_or_bytes, (str, Path)):
        file_extension = Path(file_path_or_bytes).suffix
    elif file_extension is None:
        raise ValueError("file_extension must be provided when using bytes input")
    
    parser = DocumentParserFactory.get_parser(file_extension)
    return parser.parse(file_path_or_bytes)